"""
合同文件文本抽取服务
Text extraction service for contract files (PDF / .docx / .doc)
"""

import io
import subprocess
import tempfile
import os
from typing import Tuple


class TextExtractionError(Exception):
    """统一抽取异常，触发 R3.13 file_extraction_failed"""
    pass


class TextExtractor:
    """从合同文件中抽取纯文本"""

    MAX_LENGTH = 100_000  # 文本截断上限（字符数）

    def extract(self, *, file_data: bytes, mime_type: str) -> Tuple[str, bool]:
        """
        从文件字节流中抽取纯文本。

        Args:
            file_data: 文件字节内容
            mime_type: 文件 MIME 类型

        Returns:
            (extracted_text, text_truncated)
            - extracted_text: 抽取出的纯文本（已 strip）
            - text_truncated: 是否因超过 MAX_LENGTH 而截断

        Raises:
            TextExtractionError: 文件损坏、加密、不可解析等情况
        """
        if mime_type == "application/pdf":
            text = self._extract_pdf(file_data)
        elif mime_type == (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ):
            text = self._extract_docx(file_data)
        elif mime_type == "application/msword":
            text = self._extract_doc(file_data)
        else:
            raise TextExtractionError(f"不支持的 MIME 类型: {mime_type}")

        truncated = False
        if len(text) > self.MAX_LENGTH:
            text = text[: self.MAX_LENGTH]
            truncated = True

        return text, truncated

    def _extract_pdf(self, file_data: bytes) -> str:
        """使用 pdfplumber 抽取 PDF 文本"""
        try:
            import pdfplumber
        except ImportError:
            raise TextExtractionError("pdfplumber 未安装，无法处理 PDF 文件")

        try:
            with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                # 检测加密 PDF
                if pdf.doc.is_encrypted:
                    raise TextExtractionError("PDF 文件已加密，无法抽取文本")

                pages_text = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)

                    # 提前截断：避免处理超大 PDF 时内存溢出
                    current_total = sum(len(t) for t in pages_text)
                    if current_total >= self.MAX_LENGTH * 2:
                        break

                return "\n".join(pages_text).strip()

        except TextExtractionError:
            raise
        except Exception as e:
            raise TextExtractionError(f"PDF 文本抽取失败: {str(e)}")

    def _extract_docx(self, file_data: bytes) -> str:
        """使用 python-docx 抽取 .docx 文本"""
        try:
            from docx import Document
        except ImportError:
            raise TextExtractionError("python-docx 未安装，无法处理 .docx 文件")

        MAX_EXTRACT_CHARS = 100_000  # 与 MAX_LENGTH 保持一致

        try:
            doc = Document(io.BytesIO(file_data))
            parts = []
            total_chars = 0

            # 段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    text = para.text.strip()
                    parts.append(text)
                    total_chars += len(text)
                    if total_chars >= MAX_EXTRACT_CHARS:
                        break

            # 表格文本（提前截断避免处理超大文档）
            if total_chars < MAX_EXTRACT_CHARS:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text = cell.text.strip()
                                parts.append(text)
                                total_chars += len(text)
                                if total_chars >= MAX_EXTRACT_CHARS:
                                    break
                        if total_chars >= MAX_EXTRACT_CHARS:
                            break
                    if total_chars >= MAX_EXTRACT_CHARS:
                        break
                if total_chars >= MAX_EXTRACT_CHARS:
                    pass

            return "\n".join(parts).strip()

        except Exception as e:
            raise TextExtractionError(f".docx 文本抽取失败: {str(e)}")

    def _extract_doc(self, file_data: bytes) -> str:
        """使用 antiword 抽取 .doc 文本（需要系统安装 antiword）"""
        tmp_path = None
        try:
            # 写入临时文件
            with tempfile.NamedTemporaryFile(
                suffix=".doc", delete=False
            ) as tmp_file:
                tmp_file.write(file_data)
                tmp_path = tmp_file.name

            result = subprocess.run(
                ["antiword", tmp_path],
                capture_output=True,
                text=True,
                timeout=30,
            )

            if result.returncode != 0:
                # antiword 失败，尝试 docx2txt 作为备选
                try:
                    import docx2txt
                    text = docx2txt.process(tmp_path)
                    return (text or "").strip()
                except Exception:
                    pass
                raise TextExtractionError(
                    f".doc 文本抽取失败（antiword 返回码 {result.returncode}）: "
                    f"{result.stderr}"
                )

            return result.stdout.strip()

        except TextExtractionError:
            raise
        except subprocess.TimeoutExpired:
            raise TextExtractionError(".doc 文本抽取超时")
        except FileNotFoundError:
            # antiword 未安装，尝试 docx2txt
            try:
                import docx2txt
                if tmp_path and os.path.exists(tmp_path):
                    text = docx2txt.process(tmp_path)
                    return (text or "").strip()
            except Exception:
                pass
            raise TextExtractionError(
                "antiword 未安装，无法处理 .doc 文件。"
                "请在 Dockerfile 中添加 apt-get install -y antiword"
            )
        except Exception as e:
            raise TextExtractionError(f".doc 文本抽取失败: {str(e)}")
        finally:
            # 清理临时文件
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
