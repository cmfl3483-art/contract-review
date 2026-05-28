"""
TextExtractor 单元测试
Tests for backend/app/services/text_extractor.py

Requirements: 3.13, 3.14
"""

import io
import struct
import pytest

from app.services.text_extractor import TextExtractor, TextExtractionError


# ─────────────────────────────────────────────
# 辅助函数：生成最小可读样本
# ─────────────────────────────────────────────


def make_minimal_pdf(text: str = "Hello PDF") -> bytes:
    """生成一个包含可读文本层的最小 PDF 字节流"""
    # 使用 reportlab 或手工构造最小 PDF
    # 这里手工构造一个最小 PDF（不依赖 reportlab）
    content_stream = f"BT /F1 12 Tf 100 700 Td ({text}) Tj ET"
    content_length = len(content_stream)

    pdf = f"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj

2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj

3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj

4 0 obj
<< /Length {content_length} >>
stream
{content_stream}
endstream
endobj

5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj

xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000360 00000 n 

trailer
<< /Size 6 /Root 1 0 R >>
startxref
441
%%EOF"""
    return pdf.encode("latin-1")


def make_minimal_docx(text: str = "Hello DOCX") -> bytes:
    """生成一个包含可读文本的最小 .docx 字节流"""
    try:
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph(text)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        pytest.skip("python-docx 未安装，跳过 .docx 测试")


def make_large_text_pdf(char_count: int = 110_000) -> bytes:
    """生成一个文本层超过 100000 字符的 PDF"""
    # 用 pdfplumber 能读取的方式生成大文本 PDF
    try:
        import pdfplumber
        import reportlab.lib.pagesizes as pagesizes
        from reportlab.pdfgen import canvas
        import io

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=pagesizes.A4)
        # 写入超过 100000 字符的文本（分多行）
        text = "A" * char_count
        # 每行 80 字符
        y = 800
        for i in range(0, min(len(text), 5000), 80):
            c.drawString(10, y, text[i : i + 80])
            y -= 15
            if y < 50:
                c.showPage()
                y = 800
        c.save()
        return buf.getvalue()
    except ImportError:
        pytest.skip("reportlab 未安装，跳过大文本 PDF 测试")


# ─────────────────────────────────────────────
# 测试：PDF 抽取
# ─────────────────────────────────────────────


class TestPDFExtraction:
    """PDF 文本抽取测试"""

    def test_extract_pdf_success(self):
        """正常 PDF 抽取成功，不截断"""
        try:
            import pdfplumber
        except ImportError:
            pytest.skip("pdfplumber 未安装")

        # 使用 reportlab 生成最小 PDF
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            import io

            buf = io.BytesIO()
            c = canvas.Canvas(buf, pagesize=A4)
            c.drawString(100, 700, "Hello PDF Test")
            c.save()
            pdf_bytes = buf.getvalue()
        except ImportError:
            pytest.skip("reportlab 未安装，跳过 PDF 生成测试")

        extractor = TextExtractor()
        text, truncated = extractor.extract(
            file_data=pdf_bytes, mime_type="application/pdf"
        )

        assert "Hello PDF Test" in text
        assert truncated is False

    def test_extract_corrupted_pdf_raises(self):
        """损坏的 PDF 抛出 TextExtractionError"""
        try:
            import pdfplumber
        except ImportError:
            pytest.skip("pdfplumber 未安装")

        extractor = TextExtractor()
        corrupted_bytes = b"This is not a valid PDF file at all"

        with pytest.raises(TextExtractionError):
            extractor.extract(
                file_data=corrupted_bytes, mime_type="application/pdf"
            )

    def test_extract_empty_text_pdf_returns_empty_string(self):
        """纯图片 PDF（无文本层）返回空字符串（不抛异常）"""
        try:
            import pdfplumber
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            import io

            # 生成一个没有文本的 PDF（只有空白页）
            buf = io.BytesIO()
            c = canvas.Canvas(buf, pagesize=A4)
            c.showPage()  # 空白页
            c.save()
            pdf_bytes = buf.getvalue()
        except ImportError:
            pytest.skip("reportlab 未安装")

        extractor = TextExtractor()
        text, truncated = extractor.extract(
            file_data=pdf_bytes, mime_type="application/pdf"
        )

        # 空文本层 PDF 返回空字符串（由调用方判断 empty_extracted_text）
        assert text == "" or text.strip() == ""
        assert truncated is False


# ─────────────────────────────────────────────
# 测试：DOCX 抽取
# ─────────────────────────────────────────────


class TestDocxExtraction:
    """DOCX 文本抽取测试"""

    def test_extract_docx_success(self):
        """正常 .docx 抽取成功，不截断"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")

        import io

        doc = Document()
        doc.add_paragraph("Hello DOCX Test")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        extractor = TextExtractor()
        text, truncated = extractor.extract(
            file_data=docx_bytes,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        assert "Hello DOCX Test" in text
        assert "Second paragraph" in text
        assert truncated is False

    def test_extract_corrupted_docx_raises(self):
        """损坏的 .docx 抛出 TextExtractionError"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")

        extractor = TextExtractor()
        corrupted_bytes = b"This is not a valid docx file"

        with pytest.raises(TextExtractionError):
            extractor.extract(
                file_data=corrupted_bytes,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


# ─────────────────────────────────────────────
# 测试：截断逻辑
# ─────────────────────────────────────────────


class TestTruncation:
    """文本截断测试"""

    def test_truncation_at_100000_chars(self):
        """超过 100000 字符时截断，返回 text_truncated=True 且长度恰为 100000"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")

        import io

        # 生成超过 100000 字符的 .docx
        doc = Document()
        # 每段 1000 字符，共 110 段 = 110000 字符
        for i in range(110):
            doc.add_paragraph("A" * 1000)
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        extractor = TextExtractor()
        text, truncated = extractor.extract(
            file_data=docx_bytes,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        assert truncated is True
        assert len(text) == TextExtractor.MAX_LENGTH

    def test_no_truncation_under_limit(self):
        """不超过 100000 字符时不截断"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")

        import io

        doc = Document()
        doc.add_paragraph("Short text")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        extractor = TextExtractor()
        text, truncated = extractor.extract(
            file_data=docx_bytes,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        assert truncated is False
        assert len(text) <= TextExtractor.MAX_LENGTH


# ─────────────────────────────────────────────
# 测试：不支持的 MIME 类型
# ─────────────────────────────────────────────


class TestUnsupportedMime:
    """不支持的 MIME 类型测试"""

    def test_unsupported_mime_raises(self):
        """不支持的 MIME 类型抛出 TextExtractionError"""
        extractor = TextExtractor()

        with pytest.raises(TextExtractionError, match="不支持的 MIME 类型"):
            extractor.extract(
                file_data=b"some data",
                mime_type="application/octet-stream",
            )
